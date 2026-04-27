"""Word Export Skill - Word文档导出技能"""

import io
import logging
from typing import Dict, Any, List, Optional

from app.agent.skills.core.base_skill import BaseSkill, SkillResult

logger = logging.getLogger(__name__)

# 检查 python-docx 是否可用
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx 未安装，Word导出功能不可用")


class WordExportSkill(BaseSkill):
    """Word导出 Skill - 创建带样式的Word文档"""

    DEFAULT_PARAMETERS = [
        {"name": "operation", "type": "string", "required": True, "description": "操作类型: export(导出Word), append(追加内容), read(读取Word)"},
        {"name": "content", "type": "string", "required": False, "description": "文档内容，支持Markdown格式"},
        {"name": "filename", "type": "string", "required": False, "description": "文件名（不含.docx扩展名）"},
        {"name": "title", "type": "string", "required": False, "description": "文档标题"},
        {"name": "style", "type": "string", "required": False, "description": "样式预设: normal(默认), report(报告), contract(合同), letter(信函)", "default": "normal"},
        {"name": "sections", "type": "string", "required": False, "description": "分节内容，JSON数组格式，每节包含 type 和 content"},
    ]

    def __init__(self):
        super().__init__()
        self.name = "word_export"
        self.description = "创建带样式的Word文档，或将内容导出为Word格式"
        self.parameters = self.DEFAULT_PARAMETERS

    async def execute(self, **kwargs) -> SkillResult:
        if not DOCX_AVAILABLE:
            return SkillResult(success=False, error="python-docx 库未安装，请运行: pip install python-docx")

        operation = kwargs.get("operation", "export")
        content = kwargs.get("content", "")
        filename = kwargs.get("filename", "document")
        title = kwargs.get("title", "")
        style = kwargs.get("style", "normal")
        sections = kwargs.get("sections", "")

        try:
            if operation == "export":
                return await self._export_word(content, filename, title, style)
            elif operation == "append":
                return await self._append_content(content, filename, style)
            elif operation == "read":
                return await self._read_word(content)
            else:
                return SkillResult(success=False, error=f"不支持的操作: {operation}")
        except Exception as e:
            logger.error(f"[WordExport] 操作失败: {e}")
            return SkillResult(success=False, error=str(e))

    async def _export_word(self, content: str, filename: str, title: str, style: str) -> SkillResult:
        """导出为Word文档"""

        # 创建文档
        doc = Document()

        # 应用样式
        if style == "report":
            self._apply_report_style(doc, title or filename)
        elif style == "contract":
            self._apply_contract_style(doc, title or filename)
        elif style == "letter":
            self._apply_letter_style(doc, title or filename)
        else:
            self._apply_normal_style(doc, title or filename, content)

        # 解析内容
        parsed_content = self._parse_markdown(content)

        # 添加内容段落
        for item in parsed_content:
            if item["type"] == "heading":
                doc.add_heading(item["content"], level=item.get("level", 1))
            elif item["type"] == "paragraph":
                para = doc.add_paragraph(item["content"])
                self._format_paragraph(para, item.get("format", {}))
            elif item["type"] == "list":
                for list_item in item["items"]:
                    doc.add_paragraph(list_item, style="List Bullet")
            elif item["type"] == "table":
                self._add_table(doc, item["data"])

        # 保存到字节流
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        doc_content = output.getvalue()

        # 确定文件名
        if not filename.endswith(".docx"):
            filename = f"{filename}.docx"

        # 计算段落数量
        para_count = len(parsed_content)

        return SkillResult(
            success=True,
            data={
                "filename": filename,
                "content_base64": doc_content.hex() if isinstance(doc_content, bytes) else doc_content,
                "paragraphs": para_count,
                "style": style,
            },
            metadata={
                "operation": "export",
                "paragraphs_created": para_count,
            }
        )

    async def _append_content(self, content: str, filename: str, style: str) -> SkillResult:
        """追加内容到文档"""
        return SkillResult(success=False, error="追加功能暂不支持，请使用 export 操作创建完整文档")

    async def _read_word(self, content: str) -> SkillResult:
        """读取Word文档内容"""
        return SkillResult(success=False, error="读取Word功能暂不支持")

    def _apply_normal_style(self, doc: Document, title: str, content: str) -> None:
        """应用普通样式"""
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_report_style(self, doc: Document, title: str) -> None:
        """应用报告样式"""
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # 添加副标题
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run("报告文档")
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(128, 128, 128)

    def _apply_contract_style(self, doc: Document, title: str) -> None:
        """应用合同样式"""
        if title:
            heading = doc.add_heading(title, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _apply_letter_style(self, doc: Document, title: str) -> None:
        """应用信函样式"""
        pass

    def _parse_markdown(self, content: str) -> List[Dict]:
        """解析Markdown内容"""
        items = []
        lines = content.split("\n")

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 标题
            if line.startswith("### "):
                items.append({"type": "heading", "level": 3, "content": line[4:]})
            elif line.startswith("## "):
                items.append({"type": "heading", "level": 2, "content": line[3:]})
            elif line.startswith("# "):
                items.append({"type": "heading", "level": 1, "content": line[2:]})
            # 列表
            elif line.startswith("- ") or line.startswith("* "):
                if not items or items[-1].get("type") != "list":
                    items.append({"type": "list", "items": []})
                items[-1]["items"].append(line[2:])
            # 表格
            elif "|" in line:
                if not items or items[-1].get("type") != "table":
                    items.append({"type": "table", "data": []})
                cells = [c.strip() for c in line.split("|") if c.strip()]
                if cells and not all(c == "---" for c in cells):
                    items[-1]["data"].append(cells)
            # 段落
            else:
                items.append({
                    "type": "paragraph",
                    "content": line,
                    "format": {}
                })

        return items

    def _format_paragraph(self, para, format_dict: Dict) -> None:
        """格式化段落"""
        para_runs = para.runs
        if not para_runs:
            return

        run = para_runs[0]
        if format_dict.get("bold"):
            run.font.bold = True
        if format_dict.get("italic"):
            run.font.italic = True
        if format_dict.get("color"):
            try:
                run.font.color.rgb = RGBColor(*format_dict["color"])
            except:
                pass

    def _add_table(self, doc: Document, data: List[List]) -> None:
        """添加表格"""
        if not data:
            return

        table = doc.add_table(rows=1, cols=len(data[0]))
        for i, header in enumerate(data[0]):
            table.rows[0].cells[i].text = header

        for row_data in data[1:]:
            row = table.add_row()
            for i, cell_data in enumerate(row_data):
                if i < len(row.cells):
                    row.cells[i].text = cell_data


# 导出执行入口
skill = WordExportSkill()


async def execute(**kwargs) -> SkillResult:
    """Skill 执行入口"""
    return await skill.execute(**kwargs)